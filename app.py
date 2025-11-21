import streamlit as st
import pdfplumber
import pandas as pd
import re
from io import BytesIO
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import time
from supabase import create_client, Client
import requests
from zoneinfo import ZoneInfo
from docxtpl import DocxTemplate
from docx import Document
import os
import csv
from docxcompose.composer import Composer
import tempfile
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64

# Jaunas funkcijas sākums
def create_excel_template():
    columns = [
        "Vārds uzvārds", "Uzņēmums", "Valsts kods (XX)", "Adrese 1",
        "Adrese 2", "Pasta indekss", "Grupas", "E-pasts", "Tālrunis",
        "Sūtījuma klase", "Sūtījuma tips", "Sūtījuma veids", "Apdrošināšana, €",
        "Pēcmaksa, €", "Sūtījuma svars", "Sūtījuma saturs", "Nosaukums",
        "Daudzums", "Neto svars (kg)", "Vērtība, €", "HS tarifa Nr.",
        "Izcelsmes valsts", "Papildus pakalpojumi", "Komerciāla prece",
        "Piezīmes", "AP", "MD", "PVN Nr./ Eksportētāja kods",
        "PVN Nr./ Importētāja kods", "Postage Paid", "Saistītie dokumenti",
        "Dokumenta apraksts", "Dokumenta numurs", "Adrese"  # Pievienots "Adrese"
    ]
    return pd.DataFrame(columns=columns)

def clean_address_for_Adrese2(address):
    if not isinstance(address, str):
        return ""
    idx = address.find(',')
    if idx != -1:
        result = address[:idx].strip()
    else:
        result = address.strip()
    # Nodrošina, ka pirms vārda "iela" vienmēr ir atstarpe
    result = re.sub(r'(?i)(?<!\s)(iela)', r' iela', result)
    # Nodrošinām, ka pirms un pēc "-" vienmēr ir viena atstarpe
    result = re.sub(r'\s*-\s*', ' - ', result)
    # Pēc tam pārvēršam gadījumus, kad burts "k" (neatkarīgi no lieluma) ir tieši pirms "-" – rezultātā tiks saglabāta forma "k-"
    result = re.sub(r'(?i)(k)\s*-\s*', r'\1-', result)
    
    return result

def extract_pasta_indekss(address):
    if not isinstance(address, str):
        return ""
    # Atrodam pēdējo komatu visā adresē
    idx = address.rfind(',')
    if idx == -1:
        # Ja nav komata, atgriežam pēdējo rindu
        lines = [l.strip() for l in address.split('\n') if l.strip()]
        value = lines[-1] if lines else ""
    else:
        # Atgriežam visu pēc pēdējā komata
        value = address[idx+1:].strip()
    # Noņem liekās atstarpes ap mīnus zīmi, lai rezultāts būtu, piemēram, "LV-3001"
    value = re.sub(r'\s*-\s*', '-', value)
    return value
    
def extract_valsts_kods_from_pasta_indekss(pasta_indekss):
    if not isinstance(pasta_indekss, str):
        return ""
    # Izmanto regex, lai izvilktu burtu kombināciju pirms defises
    match = re.match(r'^([A-Za-z]+)-', pasta_indekss)
    if match:
        return match.group(1)
    return ""

def extract_second_part(address):
    if not isinstance(address, str):
        return ""
    parts = address.split(',')
    
    # Izvelkam adreses daļu no pirmā līdz trešajam komatam
    if len(parts) >= 4:
        result = ', '.join(parts[1:4])
    elif len(parts) >= 2:
        result = ', '.join(parts[1:])
    else:
        return ""
    
    # Notīrām pasta indeksu un liekas atstarpes
    result = result.strip()
    
    # Noņemam pasta indeksu formātā "LV-XXXX" vai "LV- XXXX"
    result = re.sub(r',?\s*LV-\s*\d{4}(?=\s*$|\s*,)', '', result)
    
    # Ja tekstā nav "nov", noņemam visus komatus un aizstājam ar atstarpi
    if 'nov' not in result.lower():
        result = result.replace(',', ' ')
        # Aizvietojam vairākas atstarpes ar vienu
        result = ' '.join(result.split())
    else:
        # Ja ir "nov", tad tikai notīrām liekas atstarpes ap komatiem
        result = re.sub(r'\s*,\s*', ', ', result)
    
    return result.strip()

def clean_company_name(text):
    if not isinstance(text, str):
        return text
    
    # Nomainām vairākas rindiņas ar vienu atstarpi
    text = re.sub(r'\s+', ' ', text)
    
    # Labojam nepareizi savienotus vārdus (piemēram, "ValstsValsts" -> "Valsts Valsts")
    text = re.sub(r'([a-zāčēģīķļņšūž])([A-ZĀČĒĢĪĶĻŅŠŪŽ])', r'\1 \2', text)
    
    # Notīrām liekas atstarpes ap pēdiņām
    text = re.sub(r'\s*"\s*', '"', text)
    
    # Noņemam liekās atstarpes, bet saglabājam vienu atstarpi starp vārdiem
    text = ' '.join(text.split())
    
    # Pārbaudām, vai ir pāra pēdiņu skaits
    quote_count = text.count('"')
    if quote_count == 2:
        # Atrodam pirmo un pēdējo pēdiņu indeksu
        first_quote = text.find('"')
        last_quote = text.rfind('"')
        
        # Sadalām tekstu trīs daļās: pirms pēdiņām, starp pēdiņām un pēc pēdiņām
        before_quotes = text[:first_quote].strip()
        between_quotes = text[first_quote+1:last_quote].strip()
        
        # Savienojam atpakaļ ar pareizu formatējumu
        text = f"{before_quotes} \"{between_quotes}\""
    
    return text.strip()

def process_csv_data(df_csv):
    df_excel = create_excel_template()
    if "Adrese" in df_csv.columns:
        # Izmantojam clean_address_for_Adrese2 funkciju priekš Adrese 1
        df_excel["Adrese 1"] = df_csv["Adrese"].apply(clean_address_for_Adrese2)
        df_excel["Adrese 2"] = df_csv["Adrese"].apply(extract_second_part)
        
        # Apstrādājam pārējos datus
        df_excel["Pasta indekss"] = df_csv["Adrese"].apply(extract_pasta_indekss)
        df_excel["Valsts kods (XX)"] = df_excel["Pasta indekss"].apply(extract_valsts_kods_from_pasta_indekss)
        
        # Veidojam pilno adresi no kolonnām
        df_excel["Adrese"] = df_excel["Adrese 1"].fillna('') + ', ' + df_excel["Adrese 2"].fillna('')
        df_excel["Adrese"] = df_excel["Adrese"].str.replace(r',\s*,', ',', regex=True).str.strip(', ').replace('', pd.NA)

    if "VardsUzvārdsNosaukums" in df_csv.columns:
        # Vispirms notīrām un formatējam uzņēmumu nosaukumus
        df_csv["VardsUzvārdsNosaukums"] = df_csv["VardsUzvārdsNosaukums"].apply(clean_company_name)
        
        # Izveidojam masku katram uzņēmuma veidam
        sia_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("SIA", na=False, case=False)
        sabiedriba_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("Sabiedrība ar", na=False, case=False)
        valsts_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("Valsts", na=False, case=False)
        pasvaldiba_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("Pašvaldība", na=False, case=False)
        as_mask = df_csv["VardsUzvārdsNosaukums"].str.contains(r'\bAS\b', na=False, case=False)  # \b nodrošina, ka "AS" ir atsevišķs vārds
        akciju_sab_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("Akciju sabiedrība", na=False, case=False)
        ministrija_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("ministrija", na=False, case=False)
        parvalde_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("pārvalde", na=False, case=False)
        zemnieku_mask = df_csv["VardsUzvārdsNosaukums"].str.contains("saimniecība", na=False, case=False)
        
        # Apvienojam visas maskas vienā, lai identificētu uzņēmumus
        company_mask = (sia_mask | sabiedriba_mask | valsts_mask | pasvaldiba_mask | 
                       as_mask | akciju_sab_mask | ministrija_mask | parvalde_mask |
                       zemnieku_mask)
        
        # Kopējam vērtības "Vārds uzvārds" kolonnā tikai tām rindām, kur nav uzņēmums
        df_excel.loc[~company_mask, "Vārds uzvārds"] = df_csv.loc[~company_mask, "VardsUzvārdsNosaukums"]
        
        # Kopējam vērtības "Uzņēmums" kolonnā tām rindām, kur ir uzņēmums
        df_excel.loc[company_mask, "Uzņēmums"] = df_csv.loc[company_mask, "VardsUzvārdsNosaukums"]
    
    return df_excel

def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Dati')
    output.seek(0)
    
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill
    
    wb = load_workbook(output)
    ws = wb.active
    
    # Iegūstam pirmās rindas vērtības, lai noteiktu kolonnas indeksu
    header = [cell.value for cell in ws[1]]
    try:
        col_index = header.index("Valsts kods (XX)") + 1  # openpyxl izmanto 1-bāzētu indeksu
    except ValueError:
        col_index = None
    
    if col_index is not None:
        # Izveidojam sarkano aizpildījuma stilu
        red_fill = PatternFill(fill_type="solid", fgColor="FF0000")
        # Iterējam caur datu rindām (no otrās rindas)
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            cell = row[col_index - 1]
            # Pārliecināmies, ka tiek salīdzināta tukšuma simbolu attīrīta vērtība
            cell_val = str(cell.value).strip() if cell.value is not None else ""
            if cell_val != "LV":
                for c in row:
                    c.fill = red_fill
    new_output = BytesIO()
    wb.save(new_output)
    new_output.seek(0)
    return new_output.getvalue()


def download_link(file_data, file_name, link_text):
    b64 = base64.b64encode(file_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}">{link_text}</a>'
    return href

st.set_page_config(
    page_title="Vēstuļu draugs",
    layout="centered",
    initial_sidebar_state="collapsed",
)

supabase_url = "https://uhwbflqdripatfpbbetf.supabase.co"
supabase_key = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InVod2JmbHFkcmlwYXRmcGJiZXRmIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTczMDcxODE2MywiZXhwIjoyMDQ2Mjk0MTYzfQ.78wsNZ4KBg2l6zeZ1ZknBBooe0PeLtJzRU-7eXo3WTk"

try:
    supabase: Client = create_client(supabase_url, supabase_key)
except Exception as e:
    st.error(f"Neizdevās inicializēt Supabase klientu: {e}")
    st.stop()

def show_warning(message):
    st.sidebar.warning(message)
    st.warning(message)

def show_error(message):
    st.sidebar.error(message)
    st.error(message)

def show_warning_sidebar_only(message):
    st.sidebar.warning(message)

def show_error_sidebar_only(message):
    st.sidebar.error(message)

def authenticate(username, password, supabase_client):
    try:
        response = supabase_client.table("users").select("password").eq("username", username).execute()
        data = response.data
        if not data:
            return False
        stored_password = data[0]["password"]
        return password == stored_password
    except Exception as e:
        show_error("Autentifikācijas kļūda.")
        return False

def log_user_login(username):
    try:
        riga_tz = ZoneInfo('Europe/Riga')
        current_time = datetime.now(riga_tz).isoformat()
        data = {
            "username": username,
            "App": "Vēstuļu draugs",
            "Ver": "4.5",
            "app_type": "web",
            "login_time": current_time
        }
        headers = {
            "apikey": supabase_key,
            "Authorization": f"Bearer {supabase_key}",
            "Content-Type": "application/json"
        }
        url = f"{supabase_url}/rest/v1/user_data"
        response = requests.post(url, json=data, headers=headers)
        if response.status_code in [200, 201]:
            st.session_state.login_data_logged = True
        else:
            show_error("Kļūda datu ierakstīšanā.")
    except Exception as e:
        show_error(f"Kļūda: {str(e)}")

def login():
    username = st.session_state.get('username', '').strip()
    password = st.session_state.get('password', '').strip()
    if not username or not password:
        show_error("Lūdzu, ievadiet gan lietotājvārdu, gan paroli.")
    else:
        if authenticate(username, password, supabase):
            st.session_state.logged_in = True
            st.session_state.username_logged = username
            log_user_login(username)
            st.session_state.login_success = True
        else:
            show_error("Nepareizs lietotājvārds vai parole.")

def show_login():
    st.title("Vēstuļu draugs")
    with st.form(key='login_form', clear_on_submit=True):
        username = st.text_input("Lietotājvārds", key='username')
        password = st.text_input("Parole", type="password", key='password')
        submit_button = st.form_submit_button(label="Pieslēgties", on_click=login)
    if 'login_data_logged' in st.session_state and st.session_state.login_data_logged:
        st.success("Ielogošanās dati ierakstīti")
        st.session_state.login_data_logged = False
    if 'login_success' in st.session_state and st.session_state.login_success:
        st.success("Veiksmīgi pieteicies!")
        st.session_state.login_success = False
    st.markdown("<div style='text-align: center; margin-top: 20px; color: gray;'>Copyright © 2024 METRUM</div>", unsafe_allow_html=True)

def clean_address_field(address):
    if isinstance(address, str):
        address = address.replace('\r', '\n').strip()
        address = re.sub(r',+\n', '\n', address)
        address = re.sub(r'\n,+', '\n', address)
        address = re.sub(r',{2,}', ',', address)
        address = re.sub(r'^,|,$', '', address)
        address = re.sub(r'\s*,\s*', ', ', address)
        return address
    return address

def format_date_latvian(date_obj):
    month_names = {
        1: "janvārī",
        2: "februārī",
        3: "martā",
        4: "aprīlī",
        5: "maijā",
        6: "jūnijā",
        7: "jūlijā",
        8: "augustā",
        9: "septembrī",
        10: "oktobrī",
        11: "novembrī",
        12: "decembrī"
    }
    if isinstance(date_obj, (datetime, pd.Timestamp)):
        date_obj = date_obj.date()
    # Atgriežam datumu kā teksta virkni ar konkrētu formatējumu
    return str(f"{date_obj.year}. gada {date_obj.day}. {month_names.get(date_obj.month, '')}")

# Atjauninātā restore_address_format() funkcija
def restore_address_format(address):
    if not isinstance(address, str):
        return address
    text = address
    # Pievieno atstarpi starp burtiem un cipariem (piem., "gatve12" -> "gatve 12")
    text = re.sub(r'([A-Za-zāčēģīķļņšž])(\d)', r'\1 \2', text, flags=re.IGNORECASE)
    # Ja vārdi "Marsa" un "gatve" ir sapludināti, ievieto atstarpi.
    text = re.sub(r'([A-Za-zāčēģīķļņšž]+)(gatve)', r'\1 \2', text, flags=re.IGNORECASE)
    # Ja aiz mīnus zīmes seko atstarpe, aizvieto to ar rindu pārrāvumu.
    text = re.sub(r'-\s', '-\n', text)
    # Ja pēc komata nav atstarpes, ievieto to.
    text = re.sub(r',(\S)', r', \1', text)
    return text

def detect_gender_by_name(full_name):
    """
    Nosaka dzimumu pēc vārda, ņemot vērā specifiskus gadījumus un vārdu galotnes
    """
    if not isinstance(full_name, str):
        return 'M'
    
    # Atdalām vārdu no sertifikāta numura
    first_name = full_name.split()[0] if full_name.split() else ''
    first_name_lower = first_name.lower()
    
    # Definējam specifisku vīriešu vārdu sarakstu
    male_names = {
        'ēvalds', 'valdis', 'gatis', 'kristaps', 'jānis', 'andris', 'juris',
        'māris', 'kārlis', 'aigars', 'edgars', 'normunds', 'raivis', 'oskars', 
        'gunārs', 'andrejs', 'pēteris', 'arturs', 'artūrs'
    }
    
    # Definējam specifisku sieviešu vārdu sarakstu
    female_names = {
        'linda', 'anna', 'inga', 'sandra', 'ilze', 'inese', 'dace', 'kristīne',
        'maija', 'liene', 'zane', 'līga'
    }
    
    # Vispirms pārbaudam specifiskos vārdus
    if first_name_lower in male_names:
        return 'M'
    if first_name_lower in female_names:
        return 'F'
    
    # Tipiskas sieviešu vārdu galotnes
    female_endings = ['a', 'e']
    
    # Vīriešu vārdu galotnes
    male_endings = ['s', 'is', 'us']
    
    # Vispirms pārbaudām vīriešu galotnes
    for ending in male_endings:
        if first_name_lower.endswith(ending):
            return 'M'
    
    # Tad pārbaudām sieviešu galotnes
    for ending in female_endings:
        if first_name_lower.endswith(ending):
            return 'F'
    
    # Ja neviens no nosacījumiem neatbilst, pieņemam ka tas ir vīrietis
    return 'M'

def replace_gender_specific_words(doc, is_female):
    """
    Aizvieto dzimumspecifiskos vārdus dokumenta tekstā
    """
    # Pārbaudam vai tas ir sieviešu dzimtes vārds
    if is_female:
        replacements = {
            'mērnieks': 'mērniece',
            'mērniekam': 'mērniecei',
            'mērnieka': 'mērnieces',
            'mērnieku': 'mērnieci'
        }
        
        # Veicam aizvietošanu tikai tad, ja ir sieviešu dzimte
        for paragraph in doc.paragraphs:
            for old_word, new_word in replacements.items():
                if old_word in paragraph.text:
                    inline = paragraph.runs
                    for item in inline:
                        if old_word in item.text:
                            item.text = item.text.replace(old_word, new_word)

def perform_mail_merge(template_path, records, output_dir):
    output_paths = []
    try:
        template = DocxTemplate(template_path)
    except Exception as e:
        show_error("Neizdevās ielādēt šablonu.")
        return output_paths
        
    for idx, record in enumerate(records):
        try:
            context = record.copy()
            # Apstrādājam 'Adrese' lauku:
            address = record.get('Adrese', '')
            # 1. Noņemam esošos rindu pārrāvumus, aizvietojot tos ar atstarpi
            address = address.replace('\n', ' ')
            # 2. Ievietojam rindu pārrāvumu pēc katra komata (ar jebkuru atstarpju virkni pēc komata)
            address = re.sub(r',\s*', ',\n', address)
            # 3. Nodrošinām, ka vienmēr pirms vārda "iela" (neatkarīgi no lielajiem/mazajiem burtiem) ir atstarpe
            address = re.sub(r'(?i)(?<!\s)(iela)', r' iela', address)
            # 4. Nodrošinām, ka pirms un pēc "-" visur ir atstarpe
            address = re.sub(r'\s*-\s*', ' - ', address)
            # 5. Atjaunojam "LV" formātu: gadījumos, kad pēc "LV" seko "-" un četri cipari, novēršam atstarpes
            address = re.sub(r'(?i)(LV)\s*-\s*(\d{4})', r'\1-\2', address)
            context['Adrese'] = address
            # 6. Nodrošinām, ka starp burta "k" un simbola "-" vienmēr nav atstarpe
            address = re.sub(r'(?i)(k)\s*-\s*', r'\1-', address)
            context['Adrese'] = address
            
            # Apstrādājam 'VardsUzvārdsNosaukums' lauku, lai tas tiktu attēlots vienā rindā
            if 'VardsUzvārdsNosaukums' in context:
                context['VardsUzvārdsNosaukums'] = context['VardsUzvārdsNosaukums'].replace('\n', ' ')
            else:
                context['VardsUzvārdsNosaukums'] = ''
            
            # Pārbaudām dzimumu no pilna teksta (vārds + sertifikāta nr)
            mernieks = record.get('Mērnieks_Vārds_Uzvārds', '')
            is_female = False
            
            # Pārbaudām sieviešu vārdus
            if any(name in mernieks.lower() for name in [
                'lelde', 'anita', 'gunita', 'agnese', 'ieva',
                'sandra', 'kristīne', 'gunda', 'māra'
            ]):
                is_female = True
            
            # Renderējam veidni ar kontekstu 
            template.render(context)
            output_path = os.path.join(output_dir, f"merged_document_{idx+1}.docx")
            template.save(output_path)
            
            # Aizvietojam vārdus tikai ja ir sieviete
            if is_female:
                doc = Document(output_path)
                replace_gender_specific_words(doc, True)
                doc.save(output_path)
                
            output_paths.append(output_path)
        except Exception as e:
            show_error(f"Kļūda renderējot ierakstu {idx+1}: {e}")
            continue
            
    return output_paths

def merge_word_documents(file_paths, merged_output_path):
    if not file_paths:
        show_error_sidebar_only("Nav dokumentu, kas varētu tikt apvienoti.")
        return
    try:
        master = Document(file_paths[0])
        composer = Composer(master)
        for file_path in file_paths[1:]:
            doc = Document(file_path)
            composer.append(doc)
        composer.save(merged_output_path)
        doc = Document(merged_output_path)
        settings = doc.settings
        view = settings.element.find(qn('w:view'))
        if view is not None:
            view.set(qn('w:val'), 'print')
        else:
            view = OxmlElement('w:view')
            view.set(qn('w:val'), 'print')
            settings.element.append(view)
        doc.save(merged_output_path)
        show_warning_sidebar_only(f"Apvienotais dokuments saglabāts kā: {merged_output_path}")
    except Exception as e:
        show_error_sidebar_only("Kļūda apvienojot dokumentus.")

def perform_full_mail_merge(template_path, records):
    with tempfile.TemporaryDirectory() as tmpdirname:
        output_paths = perform_mail_merge(template_path, records, tmpdirname)
        if not output_paths:
            show_error_sidebar_only("Mail merge process failed. Nav izveidoti dokumenti.")
            return None
        merged_document_path = os.path.join(tmpdirname, "apvienotais_dokuments.docx")
        merge_word_documents(output_paths, merged_document_path)
        if os.path.exists(merged_document_path):
            with open(merged_document_path, "rb") as f:
                merged_file = BytesIO(f.read())
            return merged_file
        else:
            show_error_sidebar_only("Apvienotais dokuments netika izveidots.")
            return None

def group_words_into_lines(words, y_tolerance=5):
    lines = []
    current_line = []
    current_top = None
    for word in sorted(words, key=lambda x: x['top']):
        if current_top is None:
            current_top = word['top']
            current_line.append(word)
        elif abs(word['top'] - current_top) <= y_tolerance:
            current_line.append(word)
        else:
            line_text = ' '.join([w['text'] for w in sorted(current_line, key=lambda x: x['x0'])])
            lines.append({'text': line_text, 'top': current_top})
            current_line = [word]
            current_top = word['top']
    if current_line:
        line_text = ' '.join([w['text'] for w in sorted(current_line, key=lambda x: x['x0'])])
        lines.append({'text': line_text, 'top': current_top})
    return lines

def clean_property_name(name):
    name = re.sub(r'^\W+', '', name)
    name = re.sub(r'\W+$', '', name)
    name = name.strip()
    return name

def process_pdf_app():
    st.markdown("<h1 style='text-align: center; color: #AC3356;'>Vēstuļu draugs</h1>", unsafe_allow_html=True)
    # Ielādējam uzņēmumu, vietu, novadu, mērnieku un sagatavotāju sarakstus no datu bāzes
    response_company = supabase.table("VD_uzņēmums").select("uzņēmums").execute()
    error_company = getattr(response_company, "error", None)
    if error_company is not None:
        show_error_sidebar_only("Neizdevās ielādēt uzņēmumu sarakstu no datu bāzes.")
        company_options = []
    else:
        company_options = [row["uzņēmums"] for row in response_company.data if "uzņēmums" in row]
    response_place = supabase.table("VD_vieta").select('"Sagatavošanas vieta"').execute()
    error_place = getattr(response_place, "error", None)
    if error_place is not None:
        show_error_sidebar_only("Neizdevās ielādēt vietu sarakstu no datu bāzes.")
        place_options = []
    else:
        place_options = [row["Sagatavošanas vieta"] for row in response_place.data if "Sagatavošanas vieta" in row]
    response_municipality = supabase.table("VD_pagasts_un_novads").select("pagasts_un_novads").execute()
    error_municipality = getattr(response_municipality, "error", None)
    if error_municipality is not None:
        show_error_sidebar_only("Neizdevās ielādēt novadu sarakstu no datu bāzes.")
        municipality_options = []
    else:
        municipality_options = [row["pagasts_un_novads"] for row in response_municipality.data if "pagasts_un_novads" in row]
    response_surveyor = supabase.table("VD_mērnieks").select('"Vārds Uzvārds (sertifikāts Nr.) mērnieka tel. nr.","mērnieka Vārds Uzvārds"').execute()
    error_surveyor = getattr(response_surveyor, "error", None)
    if error_surveyor is not None:
        show_error_sidebar_only("Neizdevās ielādēt mērnieku sarakstu no datu bāzes.")
        surveyor_dict = {}
    else:
        surveyor_dict = {}
        for row in response_surveyor.data:
            full_name = row["mērnieka Vārds Uzvārds"]
            phone_info = row["Vārds Uzvārds (sertifikāts Nr.) mērnieka tel. nr."]
            display_value = phone_info
            surveyor_dict[display_value] = (full_name, phone_info)
    response_prepared_by = supabase.table("VD_sagatavotāja").select('"Vārds Uzvārds telefona nr.","sagatavoja e-pasts"').execute()
    error_prepared_by = getattr(response_prepared_by, "error", None)
    if error_prepared_by is not None:
        show_error_sidebar_only("Neizdevās ielādēt sagatavotāju sarakstu no datu bāzes.")
        prepared_by_dict = {}
    else:
        prepared_by_dict = {}
        for row in response_prepared_by.data:
            full_name_telefons = row["Vārds Uzvārds telefona nr."]
            email_info = row["sagatavoja e-pasts"]
            display_value = f"{full_name_telefons} (e-pasts: {email_info})"
            prepared_by_dict[display_value] = (full_name_telefons, email_info)
    with st.form("main_form"):
        pdf_file = st.file_uploader("Izvēlieties PDF (Lūdzu augšupielādējiet kadastra informāciju)", type=["pdf"], help="Nav izvēlēta kadastra informācija")
        company = st.selectbox("Izvēlieties uzņēmumu:", options=[""] + company_options, index=0, help="Izvēlieties uzņēmumu no saraksta")
        place = st.selectbox("Izvēlieties vēstules sagatavošanas vietu:", options=[""] + place_options, index=0, help="Izvēlieties vietu no saraksta")
        municipality = st.selectbox("Izvēlieties uzmērāmās zemes vienības pagastu un novadu:", options=[""] + municipality_options, index=0, help="Izvēlieties novadu no saraksta")
        meeting_place = st.text_input("Ievadiet tikšanās vietu un laiku:", value="")
        meeting_date = st.date_input("Ievadiet tikšanās datumu:", datetime.today())
        if surveyor_dict:
            surveyor_options = [""] + list(surveyor_dict.keys())
            selected_surveyor_key = st.selectbox("Izvēlieties mērnieku:", options=surveyor_options, index=0, help="Izvēlieties mērnieku no saraksta")
            if selected_surveyor_key != "":
                selected_surveyor_name, selected_surveyor_phone = surveyor_dict[selected_surveyor_key]
            else:
                selected_surveyor_name, selected_surveyor_phone = "", ""
        else:
            selected_surveyor_name = ""
            selected_surveyor_phone = ""
        if prepared_by_dict:
            prepared_by_options = [""] + list(prepared_by_dict.keys())
            selected_prepared_by_key = st.selectbox("Izvēlieties sagatavotāju:", options=prepared_by_options, index=0, help="Izvēlieties sagatavotāju no saraksta")
            if selected_prepared_by_key != "":
                selected_prepared_by_name_telefons, selected_prepared_by_email = prepared_by_dict[selected_prepared_by_key]
            else:
                selected_prepared_by_name_telefons = ""
                selected_prepared_by_email = ""
        else:
            selected_prepared_by_name_telefons = ""
            selected_prepared_by_email = ""
        cols = st.columns(2)
        with cols[0]:
            submitted = st.form_submit_button("Izpildīt")
    if submitted:
        # Validācijas pārbaudes
        if not meeting_place.strip():
            show_warning("Lūdzu, ievadiet tikšanās vietu un laiku.")
            st.stop()
        if company == "":
            show_warning("Lūdzu, izvēlieties uzņēmumu.")
            st.stop()
        if place == "":
            show_warning("Lūdzu, izvēlieties vēstules sagatavošanas vietu.")
            st.stop()
        if municipality == "":
            show_warning("Lūdzu, izvēlieties pagastu un novadu.")
            st.stop()
        if surveyor_dict and selected_surveyor_key == "":
            show_warning("Lūdzu, izvēlieties mērnieku.")
            st.stop()
        if prepared_by_dict and selected_prepared_by_key == "":
            show_warning("Lūdzu, izvēlieties sagatavotāju.")
            st.stop()
        if pdf_file is None:
            show_warning("Lūdzu, izvēlieties PDF failu, lai turpinātu.")
            st.stop()
        st.success("Faila augšupielāde veiksmīga! Sākam PDF apstrādi...")
        st.session_state.excel_data = None
        st.session_state.merged_file = None
        st.session_state.pdf_file_name = os.path.splitext(pdf_file.name)[0]
        st.session_state.file_date = datetime.today().strftime('%Y%m%d')
        with st.sidebar:
            st.title("Datu apstrāde")
            st.markdown("### Kadastra informācijas apstrādre")
            all_tables_df = []
            all_elements = []
            kadastra_apzimejumi = []
            property_names = []
            kadastra_count_per_page = {}
            property_count_per_page = {}
            tables_count_per_page = {}
            total_pages = 0
            current_kadastra_num = None
            current_property_name = None
            start_processing = False
            kadastra_page1 = None
            required_columns = ["Vārds uzvārds/\nnosaukums", "Adrese"]
            uzruna_keywords = ["pašvaldība", "SIA", "Sabiedrība ar ierobežotu atbildību"]
            sidebar_progress = st.sidebar.progress(0)
        main_progress = st.progress(0)
        with st.spinner("Apstrādā datus..."):
            with pdfplumber.open(pdf_file) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, start=1):
                    table_objects = page.find_tables()
                    tables_info = []
                    for table in table_objects:
                        bbox = table.bbox
                        tables_info.append({
                            'type': 'table',
                            'page': page_num,
                            'y0': bbox[1],
                            'data': table.extract()
                        })
                    words = page.extract_words()
                    lines = group_words_into_lines(words)
                    text_info = []
                    for line in lines:
                        text_info.append({
                            'type': 'text',
                            'page': page_num,
                            'y0': line['top'],
                            'data': line['text']
                        })
                    elements = tables_info + text_info
                    elements_sorted = sorted(elements, key=lambda x: (x['page'], x['y0']))
                    for element in elements_sorted:
                        if not start_processing:
                            if element['type'] == 'text':
                                line = element['data']
                                match_property_name = re.match(r"^3\.1\.Nekustamā īpašuma nosaukums:\s*(.*)$", line)
                                if match_property_name:
                                    property_name = match_property_name.group(1).strip()
                                    property_name = clean_property_name(property_name)
                                    property_names.append(property_name)
                                    current_property_name = property_name
                                    all_elements.append({
                                        'Lapa': page_num,
                                        'Veids': 'Nekustamā īpašuma nosaukums',
                                        'Dati': property_name
                                    })
                                    st.sidebar.markdown(f"**Atrasts '3.1. Nekustamā īpašuma nosaukums' lapā {page_num}:** {property_name}")
                                    property_count_per_page[page_num] = property_count_per_page.get(page_num, 0) + 1
                                match_kadastra = re.match(r"^1\.1\.Zemes vienības kadastra apzīmējums:\s*(.*)$", line)
                                if match_kadastra:
                                    kadastra_numurs = match_kadastra.group(1).strip()
                                    kadastra_numurs = clean_property_name(kadastra_numurs)
                                    kadastra_apzimejumi.append(kadastra_numurs)
                                    current_kadastra_num = kadastra_numurs
                                    all_elements.append({
                                        'Lapa': page_num,
                                        'Veids': 'Zemes vienības kadastra apzīmējums',
                                        'Dati': kadastra_numurs
                                    })
                                    st.sidebar.markdown(f"**Atrasts '1.1. Zemes vienības kadastra apzīmējums' lapā {page_num}:** {kadastra_numurs}")
                                    kadastra_count_per_page[page_num] = kadastra_count_per_page.get(page_num, 0) + 1
                                    if page_num == 1 and not kadastra_page1:
                                        kadastra_page1 = kadastra_numurs
                                    if "Dati par pierobežniekiem" in line or "Dati par zemes vienībām" in line:
                                        start_processing = True
                                        st.sidebar.markdown(f"## Sāk apstrādi '{line}' lapā {page_num}")
                                continue
                        if element['type'] == 'text':
                            line = element['data']
                            match = re.match(r"^1\.1\.Zemes vienības kadastra apzīmējums:.*?(\d+)$", line)
                            if match:
                                kadastra_numurs = match.group(1)
                                kadastra_numurs = clean_property_name(kadastra_numurs)
                                kadastra_apzimejumi.append(kadastra_numurs)
                                current_kadastra_num = kadastra_numurs
                                st.sidebar.markdown(f"**Atrasts kadastra apzīmējums lapā {page_num}:** {kadastra_numurs}")
                                kadastra_count_per_page[page_num] = kadastra_count_per_page.get(page_num, 0) + 1
                                all_elements.append({
                                    'Lapa': page_num,
                                    'Veids': 'Teksts',
                                    'Dati': line
                                })
                        elif element['type'] == 'table':
                            table_data = element['data']
                            df = pd.DataFrame(table_data)
                            df.dropna(axis=0, how='all', inplace=True)
                            df.dropna(axis=1, how='all', inplace=True)
                            if df.empty or len(df) < 1:
                                continue
                            header_row = df.iloc[0].astype(str).str.contains('NPK', case=False, regex=False)
                            if header_row.any():
                                df.columns = df.iloc[0]
                                df = df[1:].reset_index(drop=True)
                                df.dropna(axis=0, how='all', inplace=True)
                                df.dropna(axis=1, how='all', inplace=True)
                                if df.empty:
                                    continue
                                # Meklējam tabulā "Vārds uzvārds/\nnosaukums" kolonnu
                                if "Vārds uzvārds/\nnosaukums" in df.columns:
                                    # Notīrām un formatējam uzņēmumu nosaukumus
                                    df["Vārds uzvārds/\nnosaukums"] = df["Vārds uzvārds/\nnosaukums"].apply(clean_company_name)
                                existing_columns = [col for col in required_columns if col in df.columns]
                                missing_columns = [col for col in required_columns if col not in df.columns]
                                if missing_columns:
                                    show_warning_sidebar_only(f"Lapas {page_num}: Trūkst kolonnas {', '.join(missing_columns)}.")
                                df = df[existing_columns]
                                if current_kadastra_num:
                                    df['Kadastra Apzīmējums'] = current_kadastra_num
                                else:
                                    df['Kadastra Apzīmējums'] = None
                                df['Nekustamā īpašuma nosaukums'] = current_property_name
                                df['Lapa'] = page_num
                                all_tables_df.append(df)
                                all_elements.append({
                                    'Lapa': page_num,
                                    'Veids': 'Tabula',
                                    'Dati': df.to_dict(orient='records')
                                })
                                tables_count_per_page[page_num] = tables_count_per_page.get(page_num, 0) + 1
                                st.sidebar.markdown(f"### Lapa {page_num} - Tabula")
                                st.sidebar.dataframe(df)
                    if total_pages > 0:
                        progress_percent = int((page_num / total_pages) * 100)
                        sidebar_progress.progress(progress_percent)
                        main_progress.progress(progress_percent)
        if total_pages > 0:
            sidebar_progress.progress(100)
            main_progress.progress(100)
        for page_num in range(1, total_pages + 1):
            if page_num not in tables_count_per_page:
                tables_count_per_page[page_num] = 0
                show_warning_sidebar_only(f"Tabulas netika atrastas lapā {page_num}.")
        if all_tables_df:
            df_all = pd.concat(all_tables_df, ignore_index=True)
            # Grupējam datus un saglabājam arī lapas numuru (minimālo vērtību)
            grouped_df = df_all.groupby(["Vārds uzvārds/\nnosaukums", "Adrese"], as_index=False).agg({
                'Kadastra Apzīmējums': lambda x: ', '.join(x.dropna().unique()),
                'Nekustamā īpašuma nosaukums': 'first',
                'Lapa': 'min'
            })
            grouped_df = grouped_df.rename(columns={
                'Kadastra Apzīmējums': 'kadapz',
                "Vārds uzvārds/\nnosaukums": "VardsUzvārdsNosaukums",
                'Nekustamā īpašuma nosaukums': 'NekustamaIpaIumaNosaukums'
            })
            # Šeit izmantojam restore_address_format(), lai saglabātu sākotnējo adreses formatējumu
            grouped_df['Adrese'] = grouped_df['Adrese'].apply(restore_address_format)
            required_columns_grouped = ["VardsUzvārdsNosaukums", "Adrese"]
            missing_columns = [col for col in required_columns_grouped if col not in grouped_df.columns]
            if missing_columns:
                show_error_sidebar_only(f"Trūkst nepieciešamajām kolonnām: {', '.join(missing_columns)}")
                return
            grouped_df = grouped_df.sort_values(by='Lapa', ascending=True)
            def determine_uzruna(name):
                name_lower = name.lower()
                if any(keyword.lower() in name_lower for keyword in uzruna_keywords):
                    return "Jūsu"
                else:
                    return "Jūs vai Jūsu"
            grouped_df['uzruna'] = grouped_df["VardsUzvārdsNosaukums"].apply(determine_uzruna)
            grouped_df['NekustamaIpaIumaNosaukums'] = grouped_df['NekustamaIpaIumaNosaukums'].astype(str).apply(clean_property_name)
            if kadastra_page1:
                grouped_df['Atrasts_Zemes_Vienības_Kadastra_Apzīmējums_lapā_1'] = kadastra_page1
            else:
                grouped_df['Atrasts_Zemes_Vienības_Kadastra_Apzīmējums_lapā_1'] = "Nav atrasts"
            grouped_df['Uzņēmums'] = company
            grouped_df['Vieta'] = place
            grouped_df['Pagasts_un_Novads'] = municipality
            grouped_df['Tikšanās_vieta_un_laiks'] = meeting_place
            grouped_df['Tikšanās_datums'] = str(format_date_latvian(meeting_date))
            grouped_df['Mērnieks_Vārds_Uzvārds'] = selected_surveyor_name
            grouped_df['Mērnieks_Telefons'] = selected_surveyor_phone
            grouped_df['Sagatavotājs_Vārds_Uzvārds_Telefons'] = selected_prepared_by_name_telefons
            grouped_df['Sagatavotājs_e_pasts'] = selected_prepared_by_email
            excluded_df = grouped_df[grouped_df['VardsUzvārdsNosaukums'].str.contains(r'\(miris\)', na=False, case=False)]
            filtered_df = grouped_df[~grouped_df['VardsUzvārdsNosaukums'].str.contains(r'\(miris\)', na=False, case=False)]
            st.sidebar.markdown("## Filtrētie Adresāti (miris)")
            if not excluded_df.empty:
                st.sidebar.dataframe(excluded_df)
                excluded_csv = excluded_df.to_csv(index=False).encode('utf-8')
                st.sidebar.markdown(download_link(excluded_csv, "filtrētie_adresati_miris.csv", "Lejupielādēt filtrēto adresātu CSV failu"), unsafe_allow_html=True)
            else:
                st.sidebar.info("Nav adresātu ar '(miris)' informāciju.")
            st.sidebar.markdown("### Grupēta Tabula - Visas Lapas")
            if not grouped_df.empty:
                # Pārvēršam 'Adreses' tekstu vienā rindā, aizvietojot rindu pārrāvumus ar atstarpi
                # grouped_df['Adrese'] = grouped_df['Adrese'].str.replace('\n', ' ')
                st.sidebar.dataframe(grouped_df)
                grouped_csv = grouped_df.to_csv(index=False).encode('utf-8')
                st.sidebar.markdown(download_link(grouped_csv, "grupeta_tabula_visas_lapas.csv", "Lejupielādēt grupēto tabulu CSV failā"), unsafe_allow_html=True)
                # Izmantojam process_csv_data(), lai sagatavotu pasta sarakstu
                df_excel = process_csv_data(filtered_df)
                def remove_line_breaks(text):
                    if isinstance(text, str):
                        return text.replace('\n', ' ')
                    return text
                df_excel = df_excel.applymap(remove_line_breaks)
                st.sidebar.success("Dati veiksmīgi apstrādāti un pievienoti Excel veidnei!")
                st.sidebar.write("### Pasta saraksts")
                st.sidebar.dataframe(df_excel)
                st.sidebar.write("Lejupielādēt pasta sarakstu Excel failu")
                excel_data = to_excel(df_excel)
                st.sidebar.markdown(download_link(excel_data, 'pasta_saraksts.xlsx', "Lejupielādēt pasta sarakstu"), unsafe_allow_html=True)
                st.session_state.excel_data = excel_data
            records = filtered_df.to_dict(orient='records')
            if records:
                template_path = "template.docx"
                if not os.path.exists(template_path):
                    show_error_sidebar_only(f"Word šablona fails '{template_path}' nav atrasts.")
                else:
                    try:
                        merged_file = perform_full_mail_merge(template_path, records)
                        if merged_file:
                            st.session_state.merged_file = merged_file
                            st.session_state.file_date = meeting_date.strftime('%Y%m%d')
                        else:
                            show_error_sidebar_only("Mail merge process neizdevās.")
                    except Exception as e:
                        show_error_sidebar_only(f"Kļūda mail merge procesā: {e}")
            else:
                show_warning_sidebar_only("Nav ierakstu, kas varētu tikt izmantoti Mail Merge procesā.")

def main():
    if 'reset' in st.session_state and st.session_state.reset:
        st.session_state.reset = False
        st.experimental_rerun()
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'username_logged' not in st.session_state:
        st.session_state.username_logged = ''
    if 'login_data_logged' not in st.session_state:
        st.session_state.login_data_logged = False
    if 'login_success' not in st.session_state:
        st.session_state.login_success = False
    if 'merged_file' not in st.session_state:
        st.session_state.merged_file = None
    if 'excel_data' not in st.session_state:
        st.session_state.excel_data = None
    if 'pdf_file_name' not in st.session_state:
        st.session_state.pdf_file_name = ''
    if 'file_date' not in st.session_state:
        st.session_state.file_date = ''
    if 'company' not in st.session_state:
        st.session_state.company = ""
    if 'place' not in st.session_state:
        st.session_state.place = ""
    if 'municipality' not in st.session_state:
        st.session_state.municipality = ""
    if 'meeting_place' not in st.session_state:
        st.session_state.meeting_place = ""
    if 'meeting_date_input' not in st.session_state:
        st.session_state.meeting_date_input = datetime.today()
    if 'selected_surveyor_key' not in st.session_state:
        st.session_state.selected_surveyor_key = ""
    if 'selected_prepared_by_key' not in st.session_state:
        st.session_state.selected_prepared_by_key = ""
    if not st.session_state.logged_in:
        show_login()
    else:
        process_pdf_app()
        if st.session_state.excel_data and st.session_state.pdf_file_name and st.session_state.file_date:
            excel_file_name = f"{st.session_state.pdf_file_name}_pasta_saraksts_{st.session_state.file_date}.xlsx"
            st.markdown(download_link(st.session_state.excel_data, excel_file_name, "Lejupielādēt pasta sarakstu Excel failu"), unsafe_allow_html=True)
        elif st.session_state.excel_data:
            st.markdown(download_link(st.session_state.excel_data, 'apstrādāta_veidne.xlsx', "Lejupielādēt pasta sarakstu Excel failu"), unsafe_allow_html=True)
        if st.session_state.merged_file:
            if st.session_state.pdf_file_name and st.session_state.file_date:
                download_file_name = f"{st.session_state.pdf_file_name}_pierobežnieku_vēstules_{st.session_state.file_date}.docx"
            else:
                download_file_name = "Vēstules.docx"
            docx_b64 = base64.b64encode(st.session_state.merged_file.getvalue()).decode()
            docx_href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_b64}" download="{download_file_name}">Lejupielādēt pierobežnieku vēstules Word failu</a>'
            st.markdown(docx_href, unsafe_allow_html=True)
        if st.session_state.merged_file or st.session_state.excel_data:
            st.success("Process ir pabeigts!")
        st.markdown("<div style='text-align: center; margin-top: 20px; color: gray;'>Copyright © 2024 METRUM</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
